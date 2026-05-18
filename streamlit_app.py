import streamlit as st
import requests
import json
import pandas as pd
import time
import io
import csv

API_BASE = "http://127.0.0.1:8000"


# ── EXPORT HELPERS ─────────────────────────────────────────────────────────────

def _build_docx(report: dict, turns: list, consultation_id: str) -> bytes:
    """Build a Word document from report dict + transcript turns."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()

    # Title
    title = doc.add_heading(f"Medical Consultation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 74, 198)

    doc.add_paragraph(f"Consultation ID: {consultation_id}").runs[0].font.color.rgb = RGBColor(67, 70, 85)
    doc.add_paragraph("")

    # Report sections
    field_map = [
        ("Patient Complaints",   "patient_complaints"),
        ("Symptoms",             "symptoms"),
        ("Doctor Observations",  "doctor_observations"),
        ("Diagnosis",            "diagnosis"),
        ("Tests Recommended",    "tests_recommended"),
        ("Treatment Plan",       "treatment_plan"),
        ("Follow-up",            "follow_up_instructions"),
        ("Important Notes",      "important_notes"),
        ("ICD-10 Suggestions",   "icd10_suggestions"),
    ]
    for label, key in field_map:
        val = report.get(key)
        if val:
            h = doc.add_heading(label, level=2)
            h.runs[0].font.color.rgb = RGBColor(0, 74, 198)
            doc.add_paragraph(str(val))
            doc.add_paragraph("")

    # Medications table
    meds = report.get("prescribed_medicines", [])
    if meds:
        h = doc.add_heading("Prescribed Medicines", level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 74, 198)
        cols = ["medication_name", "dosage", "unit", "frequency", "duration", "special_instructions"]
        headers = ["Medicine", "Dosage", "Unit", "Frequency", "Duration", "Instructions"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for med in meds:
            row = table.add_row().cells
            for i, col in enumerate(cols):
                row[i].text = str(med.get(col) or "—")
        doc.add_paragraph("")

    # SOAP note
    soap = report.get("soap_note")
    if soap:
        h = doc.add_heading("SOAP Note", level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 74, 198)
        doc.add_paragraph(soap)
        doc.add_paragraph("")

    # Transcript
    if turns:
        doc.add_page_break()
        h = doc.add_heading("Consultation Transcript", level=1)
        h.runs[0].font.color.rgb = RGBColor(0, 74, 198)
        for turn in turns:
            spk = turn.get("speaker", "Unknown")
            txt = turn.get("text", "")
            ts  = turn.get("timestamp", 0)
            p = doc.add_paragraph()
            label_run = p.add_run(f"{spk}  ·  {ts:.1f}s\n")
            label_run.bold = True
            label_run.font.color.rgb = RGBColor(0, 74, 198) if spk == "Doctor" else RGBColor(0, 98, 66)
            label_run.font.size = Pt(8)
            txt_run = p.add_run(txt)
            txt_run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_csv(report: dict, turns: list, consultation_id: str) -> bytes:
    """Build a CSV with report fields + transcript rows."""
    buf = io.StringIO()
    w = csv.writer(buf)

    # Report section
    w.writerow(["SECTION", "FIELD", "VALUE"])
    field_map = [
        ("Report", "consultation_id",      consultation_id),
        ("Report", "patient_complaints",   report.get("patient_complaints", "")),
        ("Report", "symptoms",             report.get("symptoms", "")),
        ("Report", "doctor_observations",  report.get("doctor_observations", "")),
        ("Report", "diagnosis",            report.get("diagnosis", "")),
        ("Report", "tests_recommended",    report.get("tests_recommended", "")),
        ("Report", "treatment_plan",       report.get("treatment_plan", "")),
        ("Report", "follow_up_instructions", report.get("follow_up_instructions", "")),
        ("Report", "important_notes",      report.get("important_notes", "")),
        ("Report", "icd10_suggestions",    report.get("icd10_suggestions", "")),
        ("Report", "soap_note",            report.get("soap_note", "")),
    ]
    for row in field_map:
        w.writerow(row)

    # Medicines
    meds = report.get("prescribed_medicines", [])
    if meds:
        w.writerow([])
        w.writerow(["MEDICINES", "medication_name", "dosage", "unit", "frequency", "duration", "special_instructions"])
        for med in meds:
            w.writerow([
                "Medicine",
                med.get("medication_name", ""),
                med.get("dosage", ""),
                med.get("unit", ""),
                med.get("frequency", ""),
                med.get("duration", ""),
                med.get("special_instructions", ""),
            ])

    # Transcript
    if turns:
        w.writerow([])
        w.writerow(["TRANSCRIPT", "speaker", "timestamp_s", "text"])
        for turn in turns:
            w.writerow([
                "Turn",
                turn.get("speaker", ""),
                turn.get("timestamp", ""),
                turn.get("text", ""),
            ])

    return buf.getvalue().encode("utf-8")

st.set_page_config(
    page_title="MedGemma — Prescription & Consultation",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ── GLOBAL STYLES ──────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Manrope:wght@400;600;700;800&family=Inter:wght@400;500;600&display=swap');

html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
#MainMenu, footer, header { visibility: hidden; }
.stDeployButton { display: none; }

.topnav {
    background: #f7f9fb;
    border-bottom: 1px solid #c3c6d7;
    padding: 0 48px;
    height: 64px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    position: sticky;
    top: 0;
    z-index: 999;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}
.topnav-brand {
    display: flex; align-items: center; gap: 10px;
    font-family: 'Manrope', sans-serif;
    font-size: 1.4rem; font-weight: 800;
    color: #004ac6;
}
.block-container { padding: 2rem 3rem 2rem 3rem !important; max-width: 100% !important; background: #f7f9fb; }

.page-title {
    font-family: 'Manrope', sans-serif;
    font-size: 2.1rem; font-weight: 700;
    color: #191c1e; margin-bottom: 6px;
}
.page-subtitle { color: #434655; font-size: 1rem; margin-bottom: 28px; }

.step-card {
    background: #ffffff;
    border: 1px solid #c3c6d7;
    border-radius: 12px;
    padding: 24px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.05);
    margin-bottom: 4px;
}
.step-label {
    font-size: 0.7rem; font-weight: 700; letter-spacing: 0.1em;
    text-transform: uppercase; color: #004ac6; margin-bottom: 10px;
    display: flex; align-items: center; gap: 8px;
}
.dot {
    width: 8px; height: 8px; border-radius: 50%;
    background: #004ac6; display: inline-block;
    animation: pulse 2s infinite;
}
@keyframes pulse { 0%,100%{opacity:1} 50%{opacity:0.35} }
.step-title {
    font-family: 'Manrope', sans-serif;
    font-size: 1.25rem; font-weight: 700; color: #191c1e; margin-bottom: 4px;
}
.step-desc { font-size: 0.84rem; color: #434655; margin-bottom: 16px; }

.upload-zone {
    border: 2px dashed #c3c6d7; border-radius: 10px;
    padding: 36px 20px; text-align: center; background: #f2f4f6;
}
.upload-hint { font-size: 0.75rem; color: #737686; margin-top: 8px; }

.locked-badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: #fff; border: 1px solid #c3c6d7;
    border-radius: 99px; padding: 8px 18px;
    font-size: 0.8rem; font-weight: 600; color: #434655;
    box-shadow: 0 2px 6px rgba(0,0,0,0.07);
    margin-top: 10px;
}

.sec-title {
    font-family: 'Manrope', sans-serif;
    font-size: 1.15rem; font-weight: 700;
    color: #191c1e; margin-bottom: 14px;
    display: flex; align-items: center; gap: 8px;
}
.sec-label {
    font-size: 0.7rem; font-weight: 700;
    text-transform: uppercase; letter-spacing: 0.1em;
    color: #004ac6; margin-bottom: 10px;
}

.badge { display: inline-flex; align-items: center; gap: 5px; padding: 3px 10px; border-radius: 99px; font-size: 0.72rem; font-weight: 700; }
.badge-completed { background: #e6f9f2; color: #006242; }
.badge-pending   { background: #dbe1ff; color: #004ac6; }
.badge-rejected  { background: #ffdad6; color: #ba1a1a; }

.records-table { width: 100%; border-collapse: collapse; background: #fff; border-radius: 12px; overflow: hidden; border: 1px solid #c3c6d7; box-shadow: 0 1px 4px rgba(0,0,0,0.05); }
.records-table th { background: #f2f4f6; padding: 12px 20px; font-size: 0.72rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.08em; color: #737686; text-align: left; border-bottom: 1px solid #c3c6d7; }
.records-table td { padding: 14px 20px; font-size: 0.88rem; color: #191c1e; border-bottom: 1px solid #e0e3e5; }
.records-table tr:last-child td { border-bottom: none; }
.records-table tr:hover td { background: #f2f4f6; }

/* Consultation-specific */
.rec-dot { display: inline-block; width: 10px; height: 10px; border-radius: 50%; background: #ba1a1a; animation: blink 1s infinite; margin-right: 6px; }
@keyframes blink { 0%,100%{opacity:1} 50%{opacity:0.2} }
.turn-doctor  { background: #dbe1ff; border-radius: 10px; padding: 8px 14px; margin: 4px 0; font-size: 0.88rem; }
.turn-patient { background: #e6f9f2; border-radius: 10px; padding: 8px 14px; margin: 4px 0; font-size: 0.88rem; }
.section-card { background: #fff; border: 1px solid #c3c6d7; border-radius: 12px; padding: 20px 24px; margin-bottom: 16px; }
</style>
""", unsafe_allow_html=True)

# ── TOP NAV ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="topnav">
    <div class="topnav-brand">⚕️ MedGemma</div>
    <div style="font-size:0.85rem;color:#737686;">Medical AI — Prescription Review & Consultation Scribe</div>
</div>
""", unsafe_allow_html=True)

# ── SESSION STATE ──────────────────────────────────────────────────────────────
_defaults = {
    # Prescription tab
    "extraction_data":    None,
    "review_response":    None,
    "upload_data":        None,
    "uploaded_file_name": None,
    # Consultation tab
    "consultation_id":    None,
    "recording":          False,
    "turns":              [],
    "report":             None,
    "summary":            None,
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── TABS ───────────────────────────────────────────────────────────────────────
tab_rx, tab_consult, tab_history = st.tabs([
    "💊 Prescription Review",
    "🎙️ Consultation Scribe",
    "🗃️ Patient History",
])


# ══════════════════════════════════════════════════════════════════════════════
#  TAB 1 — PRESCRIPTION REVIEW  (original, untouched logic)
# ══════════════════════════════════════════════════════════════════════════════
with tab_rx:
    st.markdown("""
    <div class="page-title">Prescription Review</div>
    <div class="page-subtitle">Upload a prescription image — AI extracts and structures all clinical data.</div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown("""
        <div class="step-card">
            <div class="step-label"><span class="dot"></span> STEP 01</div>
            <div class="step-title">1. Upload Prescription</div>
            <div class="step-desc">Choose an image of the physical prescription to initiate scanning.</div>
        </div>
        """, unsafe_allow_html=True)

        uploaded_file = st.file_uploader(
            "Upload", type=["png", "jpg", "jpeg"], label_visibility="collapsed"
        )

        if uploaded_file:
            st.image(uploaded_file, use_container_width=True)
            if st.session_state.uploaded_file_name != uploaded_file.name:
                with st.spinner("Uploading…"):
                    try:
                        resp = requests.post(
                            f"{API_BASE}/upload",
                            files={"file": (uploaded_file.name, uploaded_file, uploaded_file.type)}
                        )
                        st.session_state.upload_data        = resp.json()
                        st.session_state.uploaded_file_name = uploaded_file.name
                        st.session_state.extraction_data    = None
                        st.session_state.review_response    = None
                    except Exception as e:
                        st.error(f"Upload failed: {e}")

            if st.session_state.upload_data:
                st.success(f"**{st.session_state.upload_data.get('filename', uploaded_file.name)}** ready for extraction")
        else:
            st.markdown("""
            <div class="upload-zone">
                <div style="font-size:2.5rem;margin-bottom:12px">☁️</div>
                <div style="font-size:0.88rem;color:#434655;font-weight:500">Drop your file here or use the uploader above</div>
                <div class="upload-hint">200MB per file &nbsp;·&nbsp; PNG, JPG, JPEG</div>
            </div>
            """, unsafe_allow_html=True)

    with col2:
        unlocked = uploaded_file is not None and st.session_state.upload_data is not None
        opacity  = "1" if unlocked else "0.5"

        st.markdown(f"""
        <div class="step-card" style="opacity:{opacity}">
            <div class="step-label">STEP 02</div>
            <div class="step-title">2. Run Extraction</div>
            <div class="step-desc">AI reads the prescription and extracts structured clinical data.</div>
        </div>
        """, unsafe_allow_html=True)

        if unlocked:
            if st.button("▶ Run Analysis", type="primary", use_container_width=True):
                with st.spinner("MedGemma reading prescription… (~60s)"):
                    try:
                        r = requests.get(
                            f"{API_BASE}/extract",
                            params={"file_name": uploaded_file.name}
                        )
                        st.session_state.extraction_data = r.json()
                        st.session_state.review_response = None
                        st.success("Extraction complete!")
                    except Exception as e:
                        st.error(str(e))

            if st.session_state.extraction_data and "record_id" in st.session_state.extraction_data:
                data        = st.session_state.extraction_data
                review_info = data.get("review", {})
                record_id   = data.get("record_id")

                st.markdown("<br>", unsafe_allow_html=True)

                review_status = review_info.get("status", "pending")
                status_badge_map = {
                    "approved": ("badge-completed", "Approved"),
                    "rejected": ("badge-rejected",  "Rejected"),
                    "edited":   ("badge-pending",   "Edited"),
                    "pending":  ("badge-pending",   "Pending"),
                }
                badge_cls, badge_lbl = status_badge_map.get(review_status, ("badge-pending", "Pending"))

                st.markdown(f"""
                <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:12px;">
                    <div class="sec-title" style="margin-bottom:0">Extracted JSON</div>
                    <div style="display:flex;gap:8px;align-items:center;">
                        <span style="font-size:0.76rem;background:#dbe1ff;color:#004ac6;padding:3px 12px;border-radius:99px;font-weight:600;">{record_id}</span>
                        <span class="badge {badge_cls}">{badge_lbl}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                st.caption("Edit the JSON directly below. Changes are sent to the database on Approve / Save Edits / Reject.")

                json_key = f"json_edit_{record_id}"
                if json_key not in st.session_state:
                    st.session_state[json_key] = json.dumps(data, indent=2)

                edited_json_str = st.text_area(
                    label="json_editor",
                    label_visibility="collapsed",
                    value=st.session_state[json_key],
                    height=520,
                    key=f"textarea_{record_id}",
                )

                # Parse JSON — always attempt, never block buttons
                try:
                    edited_data = json.loads(edited_json_str)
                except json.JSONDecodeError:
                    edited_data = json.loads(st.session_state[json_key])

                st.markdown("<br>", unsafe_allow_html=True)
                reviewer_notes = st.text_area(
                    "Reviewer Notes (optional)",
                    placeholder="Add clinical notes, flags, or corrections…",
                    height=80,
                    key=f"notes_{record_id}"
                )

                b1, b2, b3, _ = st.columns([1, 1, 1, 3])

                def _submit(status: str):
                    corrected_extraction = edited_data.get("extracted", {})
                    try:
                        requests.put(
                            f"{API_BASE}/review/{record_id}",
                            params={"status": status, "reviewer_notes": reviewer_notes},
                            json={"extracted": corrected_extraction}
                        )
                        updated = edited_data.copy()
                        updated["review"]["status"] = status
                        st.session_state[json_key]       = json.dumps(updated, indent=2)
                        st.session_state.extraction_data = updated
                    except Exception as e:
                        st.error(str(e))

                with b1:
                    if st.button("Approve", type="primary", use_container_width=True):
                        _submit("approved"); st.success("Approved")
                with b2:
                    if st.button("Save Edits", use_container_width=True):
                        _submit("edited"); st.info("Edits saved")
                with b3:
                    if st.button("Reject", use_container_width=True):
                        _submit("rejected"); st.warning("Rejected")
        else:
            st.markdown('<div class="locked-badge">🔒 &nbsp;Complete Step 1 First</div>', unsafe_allow_html=True)

    # Stored records expander (original)
    st.markdown("<br><br>", unsafe_allow_html=True)
    with st.expander("📁 View Stored Records & Export", expanded=False):
        try:
            records_resp = requests.get(f"{API_BASE}/records")
            records      = records_resp.json()

            if records:
                rows = []
                for rec in records:
                    review = json.loads(rec.get("review_json", "{}"))
                    status = review.get("status", "pending")
                    badge  = {
                        "approved": '<span class="badge badge-completed">● Completed</span>',
                        "rejected": '<span class="badge badge-rejected">● Rejected</span>',
                    }.get(status, '<span class="badge badge-pending">● Pending</span>')
                    rows.append({
                        "uploaded":  rec.get("source_file", "—"),
                        "status":    badge,
                        "record_id": rec.get("record_id", "—"),
                    })

                table_rows = "".join(
                    f"<tr><td style='color:#434655;font-weight:600'>{r['uploaded']}</td><td>{r['status']}</td><td style='color:#737686;font-size:0.78rem'>{r['record_id']}</td></tr>"
                    for r in rows
                )
                st.markdown(f"""
                <table class="records-table">
                    <thead><tr><th>Uploaded File</th><th>Status</th><th>Record ID</th></tr></thead>
                    <tbody>{table_rows}</tbody>
                </table>
                <div style="margin-top:8px;font-size:0.78rem;color:#737686">{len(rows)} records total</div>
                """, unsafe_allow_html=True)

                st.markdown("<br>", unsafe_allow_html=True)
                dl1, dl2, _ = st.columns([1, 1, 4])
                with dl1:
                    st.download_button("⬇️ Export as JSON", json.dumps(records, indent=2),
                                       "medgemma_records.json", "application/json", use_container_width=True)
                with dl2:
                    df = pd.json_normalize(records)
                    st.download_button("⬇️ Export as CSV", df.to_csv(index=False).encode(),
                                       "medgemma_records.csv", "text/csv", use_container_width=True)
            else:
                st.info("No records yet.")
        except Exception as e:
            st.error(str(e))


# ══════════════════════════════════════════════════════════════════════════════
#  TAB 2 — CONSULTATION SCRIBE
# ══════════════════════════════════════════════════════════════════════════════
with tab_consult:
    st.markdown("""
    <div class="page-title">Consultation Scribe</div>
    <div class="page-subtitle">AI listens to the doctor-patient conversation, transcribes in real time, and generates a structured medical report.</div>
    """, unsafe_allow_html=True)

    left, right = st.columns([1, 1], gap="large")

    # ── LEFT PANEL ─────────────────────────────────────────────────────────────
    with left:
        st.markdown('<div class="sec-label">Session Info</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        patient_id = c1.text_input("Patient ID", placeholder="e.g. PAT001", key="consult_pid")
        doctor_id  = c2.text_input("Doctor ID",  placeholder="e.g. DR001",  key="consult_did")

        st.markdown("<br>", unsafe_allow_html=True)
        b1, b2 = st.columns(2)

        with b1:
            if st.button(
                "🔴 Start Consultation", type="primary",
                use_container_width=True,
                disabled=st.session_state.recording
            ):
                try:
                    r = requests.post(
                        f"{API_BASE}/consultation/start",
                        params={"patient_id": patient_id, "doctor_id": doctor_id}
                    )
                    data = r.json()
                    st.session_state.consultation_id = data["consultation_id"]
                    st.session_state.recording       = True
                    st.session_state.turns           = []
                    st.session_state.report          = None
                    st.session_state.summary         = None
                    requests.post(f"{API_BASE}/transcription/start_pipeline")
                    st.rerun()
                except Exception as e:
                    st.error(str(e))


        with b2:
            if st.button(
                "⏹ Stop & Generate Report",
                use_container_width=True,
                disabled=not st.session_state.recording
            ):
                try:
                    requests.post(f"{API_BASE}/consultation/stop")
                    st.session_state.recording = False
                    cid = st.session_state.consultation_id

                    # ── Check transcript has turns before generating ──
                    transcript_resp = requests.get(
                        f"{API_BASE}/transcription/{cid}/full", timeout=5
                    )
                    turns = transcript_resp.json().get("turns", [])

                    if not turns:
                        st.warning(
                            "⚠️ No speech was captured.\n\n"
                            "**Possible reasons:**\n"
                            "- Mic not allowed in browser (check address bar 🎤)\n"
                            "- Audio was too quiet / filtered as silence\n\n"
                            "Fix your mic and start a new consultation."
                        )
                    else:
                        with st.spinner(f"MedGemma generating report from {len(turns)} turn(s)…"):
                            r = requests.post(f"{API_BASE}/report/generate/{cid}")
                            if r.status_code == 200:
                                data = r.json()
                                st.session_state.report  = data.get("report", {})
                                st.session_state.summary = data.get("summary", "")
                                st.rerun()
                            else:
                                detail = r.json().get("detail", r.text)
                                st.error(f"Report generation failed: {detail}")
                except Exception as e:
                    st.error(str(e))

        st.markdown("<br>", unsafe_allow_html=True)

        # Live transcript label
        if st.session_state.recording:
            st.markdown('<div class="sec-label"><span class="rec-dot"></span>Live Transcript</div>', unsafe_allow_html=True)
        elif st.session_state.turns:
            st.markdown('<div class="sec-label">Transcript</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div style="color:#737686;font-size:0.88rem;">Start a consultation to see the live transcript here.</div>', unsafe_allow_html=True)

        # Poll for new turns while recording
        if st.session_state.recording and st.session_state.consultation_id:
            try:
                r = requests.get(
                    f"{API_BASE}/transcription/{st.session_state.consultation_id}/full",
                    timeout=3
                )
                st.session_state.turns = r.json().get("turns", [])
            except Exception:
                pass
            time.sleep(0.5)
            st.rerun()

        # Render transcript turns
        transcript_html = ""
        for turn in st.session_state.turns:
            spk   = turn.get("speaker", "Unknown")
            txt   = turn.get("text", "")
            ts    = turn.get("timestamp", 0)
            cls   = "turn-doctor"  if spk == "Doctor"  else "turn-patient"
            color = "#004ac6"      if spk == "Doctor"  else "#006242"
            transcript_html += f"""
            <div class="{cls}">
                <div class="speaker-label" style="font-size:0.7rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:{color};margin-bottom:2px;">{spk} &nbsp;·&nbsp; {ts:.1f}s</div>
                {txt}
            </div>"""

        if transcript_html:
            st.markdown(f'<div style="max-height:480px;overflow-y:auto;">{transcript_html}</div>', unsafe_allow_html=True)

    # ── RIGHT PANEL — report ───────────────────────────────────────────────────
    with right:
        st.markdown('<div class="sec-label">Structured Medical Report</div>', unsafe_allow_html=True)

        if not st.session_state.report:
            st.markdown('<div style="color:#737686;font-size:0.88rem;">Report will appear here after the consultation ends.</div>', unsafe_allow_html=True)
        else:
            report = st.session_state.report
            cid    = st.session_state.consultation_id

            # ── Summary banner ──────────────────────────────────────────────
            if st.session_state.summary:
                st.info(f"📝 {st.session_state.summary}")

            # ── Export buttons (Word + CSV) ─────────────────────────────────
            turns_for_export = st.session_state.get("turns", [])

            exp1, exp2, exp3 = st.columns([1, 1, 3])
            with exp1:
                docx_bytes = _build_docx(report, turns_for_export, cid)
                if docx_bytes:
                    st.download_button(
                        "⬇️ Word (.docx)",
                        data=docx_bytes,
                        file_name=f"report_{cid}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                else:
                    st.caption("Install `python-docx` for Word export")
            with exp2:
                csv_bytes = _build_csv(report, turns_for_export, cid)
                st.download_button(
                    "⬇️ CSV",
                    data=csv_bytes,
                    file_name=f"report_{cid}.csv",
                    mime="text/csv",
                    use_container_width=True,
                )

            st.markdown("<br>", unsafe_allow_html=True)

            # ── Report cards (styled like transcript, not JSON) ─────────────
            fields = [
                ("🤒", "Patient Complaints",  "patient_complaints"),
                ("🌡", "Symptoms",            "symptoms"),
                ("🔬", "Doctor Observations", "doctor_observations"),
                ("🏥", "Diagnosis",           "diagnosis"),
                ("🧪", "Tests Recommended",   "tests_recommended"),
                ("📋", "Treatment Plan",      "treatment_plan"),
                ("🔁", "Follow-up",           "follow_up_instructions"),
                ("⚠️", "Important Notes",     "important_notes"),
                ("🏷", "ICD-10 Suggestions",  "icd10_suggestions"),
            ]
            report_html = ""
            for icon, label, key in fields:
                val = report.get(key)
                if val:
                    report_html += f"""
                    <div style="background:#dbe1ff;border-radius:10px;padding:10px 16px;margin:6px 0;">
                        <div style="font-size:0.7rem;font-weight:700;text-transform:uppercase;
                                    letter-spacing:0.08em;color:#004ac6;margin-bottom:4px;">
                            {icon} {label}
                        </div>
                        <div style="font-size:0.88rem;color:#191c1e;">{val}</div>
                    </div>"""

            # Medicines as cards
            meds = report.get("prescribed_medicines", [])
            if meds:
                report_html += """
                <div style="background:#dbe1ff;border-radius:10px;padding:10px 16px;margin:6px 0;">
                    <div style="font-size:0.7rem;font-weight:700;text-transform:uppercase;
                                letter-spacing:0.08em;color:#004ac6;margin-bottom:8px;">
                        💊 Prescribed Medicines
                    </div>"""
                for med in meds:
                    name  = med.get("medication_name", "—")
                    dose  = med.get("dosage", "")
                    unit  = med.get("unit", "")
                    freq  = med.get("frequency", "")
                    dur   = med.get("duration", "")
                    instr = med.get("special_instructions", "")
                    detail_parts = []
                    if dose or unit: detail_parts.append(f"{dose} {unit}".strip())
                    if freq:         detail_parts.append(freq)
                    if dur:          detail_parts.append(dur)
                    if instr:        detail_parts.append(instr)
                    detail = "  ·  ".join(detail_parts) if detail_parts else ""
                    report_html += f"""
                    <div style="background:#fff;border-radius:8px;padding:8px 12px;margin:4px 0;
                                border-left:3px solid #004ac6;">
                        <div style="font-size:0.88rem;font-weight:600;color:#191c1e;">{name}</div>
                        <div style="font-size:0.78rem;color:#434655;margin-top:2px;">{detail}</div>
                    </div>"""
                report_html += "</div>"

            # SOAP note
            soap = report.get("soap_note")
            if soap:
                soap_escaped = soap.replace("\n", "<br>")
                report_html += f"""
                <div style="background:#f2f4f6;border-radius:10px;padding:10px 16px;margin:6px 0;
                            border:1px solid #c3c6d7;">
                    <div style="font-size:0.7rem;font-weight:700;text-transform:uppercase;
                                letter-spacing:0.08em;color:#004ac6;margin-bottom:6px;">
                        📄 SOAP Note
                    </div>
                    <div style="font-size:0.82rem;color:#191c1e;font-family:monospace;
                                line-height:1.6;">{soap_escaped}</div>
                </div>"""

            if report_html:
                st.markdown(
                    f'<div style="max-height:600px;overflow-y:auto;">{report_html}</div>',
                    unsafe_allow_html=True
                )

            # ── Raw JSON editor (collapsed) ─────────────────────────────────
            st.markdown("<br>", unsafe_allow_html=True)
            with st.expander("🔍 View / Edit Raw JSON", expanded=False):
                json_key = f"report_json_{cid}"
                if json_key not in st.session_state:
                    st.session_state[json_key] = json.dumps(report, indent=2)

                edited = st.text_area(
                    "report_editor", label_visibility="collapsed",
                    value=st.session_state[json_key], height=400,
                    key=f"report_ta_{cid}"
                )
                try:
                    parsed = json.loads(edited)
                except json.JSONDecodeError:
                    parsed = json.loads(st.session_state[json_key])

                if st.button("💾 Save JSON edits"):
                    try:
                        requests.put(f"{API_BASE}/report/{cid}", json={"report": parsed})
                        st.session_state.report    = parsed
                        st.session_state[json_key] = json.dumps(parsed, indent=2)
                        st.success("Saved to DB.")
                    except Exception as e:
                        st.error(str(e))

    # ── Consultation history (bottom of tab) ───────────────────────────────────
    st.markdown("<br><br>", unsafe_allow_html=True)
    with st.expander("🗃️ Consultation History", expanded=False):
        try:
            resp    = requests.get(f"{API_BASE}/consultations")
            records = resp.json()
            if records:
                rows = []
                for rec in records:
                    rows.append({
                        "ID":      rec.get("consultation_id", "—"),
                        "Patient": rec.get("patient_id") or "—",
                        "Doctor":  rec.get("doctor_id")  or "—",
                        "Started": (rec.get("started_at") or "—")[:19],
                        "Status":  rec.get("status", "—"),
                        "Summary": rec.get("summary")   or "—",
                    })
                st.dataframe(rows, use_container_width=True, hide_index=True)
            else:
                st.info("No consultations yet.")
        except Exception as e:
            st.error(str(e))


# ══════════════════════════════════════════════════════════════════════════════
#  TAB 3 — PATIENT HISTORY
# ══════════════════════════════════════════════════════════════════════════════
with tab_history:
    st.markdown("""
    <div class="page-title">Patient History</div>
    <div class="page-subtitle">Look up all consultation transcripts stored for a patient.</div>
    """, unsafe_allow_html=True)

    ph_col1, ph_col2 = st.columns([1, 3])
    with ph_col1:
        lookup_pid = st.text_input("Enter Patient ID", placeholder="e.g. PAT001", key="history_pid")
        search_btn = st.button("🔍 Search", type="primary", use_container_width=True)

    if search_btn and lookup_pid:
        try:
            r       = requests.get(f"{API_BASE}/patient/{lookup_pid}/transcripts")
            history = r.json()

            if not history:
                st.info(f"No records found for patient **{lookup_pid}**.")
            else:
                st.markdown(f"**{len(history)} consultation(s) found for {lookup_pid}**")
                st.markdown("<br>", unsafe_allow_html=True)

                for entry in history:
                    cid      = entry.get("consultation_id", "—")
                    started  = (entry.get("started_at") or "—")[:19]
                    ended    = (entry.get("ended_at")   or "—")[:19]
                    duration = entry.get("duration_seconds")
                    dur_str  = f"{duration//60}m {duration%60}s" if duration else "—"
                    doctor   = entry.get("doctor_id") or "—"
                    summary  = entry.get("summary")   or "No summary"
                    turns    = entry.get("transcript", [])

                    with st.expander(f"📋 {cid} — {started} | Dr: {doctor} | {dur_str}", expanded=False):
                        if summary != "No summary":
                            st.info(f"📝 {summary}")

                        # Transcript
                        if turns:
                            st.markdown('<div class="sec-label">Transcript</div>', unsafe_allow_html=True)
                            html = ""
                            for turn in turns:
                                spk   = turn.get("speaker", "Unknown")
                                txt   = turn.get("text", "")
                                ts    = turn.get("timestamp", 0)
                                cls   = "turn-doctor"  if spk == "Doctor"  else "turn-patient"
                                color = "#004ac6"      if spk == "Doctor"  else "#006242"
                                html += f"""
                                <div class="{cls}" style="margin:4px 0;">
                                    <span style="font-size:0.7rem;font-weight:700;color:{color};text-transform:uppercase;">{spk} · {ts:.1f}s</span><br>{txt}
                                </div>"""
                            st.markdown(f'<div style="max-height:360px;overflow-y:auto;">{html}</div>', unsafe_allow_html=True)
                        else:
                            st.caption("No transcript turns recorded.")

                        # Report fields
                        report = entry.get("report", {})
                        if report and report.get("diagnosis"):
                            st.markdown("<br>", unsafe_allow_html=True)
                            st.markdown('<div class="sec-label">Medical Report</div>', unsafe_allow_html=True)
                            for label, key in [
                                ("Diagnosis", "diagnosis"),
                                ("Treatment Plan", "treatment_plan"),
                                ("Follow-up", "follow_up_instructions"),
                            ]:
                                val = report.get(key)
                                if val:
                                    st.markdown(f"**{label}:** {val}")

                        # Download transcript
                        st.download_button(
                            "⬇️ Download Transcript JSON",
                            data=json.dumps(turns, indent=2),
                            file_name=f"transcript_{cid}.json",
                            mime="application/json",
                            key=f"dl_{cid}"
                        )
        except Exception as e:
            st.error(f"Error fetching history: {e}")
    elif search_btn:
        st.warning("Please enter a Patient ID first.")